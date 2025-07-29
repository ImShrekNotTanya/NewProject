import sys
import textwrap
import traceback

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import matplotlib.pyplot as plt
from PyQt5.QtGui import QColor, QRegExpValidator, QFont, QKeySequence, QPalette
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QTabWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QLineEdit, QPushButton, QScrollArea, QFrame, QListWidget,
                             QGroupBox, QMessageBox, QTableWidget, QTableWidgetItem, QAbstractItemView, QGridLayout,
                             QSizePolicy, QButtonGroup, QHeaderView, QRadioButton, QAction, QShortcut, QFileDialog)
from PyQt5.QtCore import Qt, QRegExp, QTimer
from backend import AHPBackend

class AHPFrontend(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Расчетный сервис МАИ")

        # Инициализация атрибутов
        self.backend = AHPBackend()
        self.current_matrix = None
        self.current_items = []
        self.result_display_mode = "chart"
        self.result_data = None
        self.matrix_entries = {}
        self.display_percent = False
        self.selected_levels = 3  # По умолчанию 3 уровня

        # Настройки темы и масштаба
        self.dark_mode = False
        self.current_scale = 1.0
        self.MIN_SCALE = 0.5
        self.MAX_SCALE = 2.0
        self.SCALE_STEP = 0.1

        # Создание виджетов
        self._create_widgets()
        self._setup_ui()
        self._create_settings_menu()

        # Применяем тему и масштаб
        self._apply_theme()
        self._apply_scale()

        # Настройка окна
        self.setWindowState(Qt.WindowMaximized)
        self.setWindowFlags(self.windowFlags() | Qt.WindowMaximizeButtonHint)

        # Принудительное отображение максимизированного окна
        QTimer.singleShot(100, self.force_maximize)

    def _setup_ui(self):
        """Настройка основного интерфейса"""
        try:
            central_widget = QWidget()
            central_widget.setLayout(QVBoxLayout())

            if hasattr(self, 'tabs'):
                central_widget.layout().addWidget(self.tabs)

            self.setCentralWidget(central_widget)

            if hasattr(self, 'tabs'):
                self.tabs.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

            central_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        except Exception as e:
            QMessageBox.critical(None, "Ошибка инициализации", f"Ошибка настройки интерфейса: {str(e)}")
            raise

    def _create_settings_menu(self):
        """Создаем меню для настроек темы и масштаба"""
        settings_menu = self.menuBar().addMenu("Настройки")

        # Меню темы
        theme_menu = settings_menu.addMenu("Тема")
        light_action = QAction("Светлая", self)
        light_action.triggered.connect(lambda: self.set_theme(False))
        dark_action = QAction("Темная", self)
        dark_action.triggered.connect(lambda: self.set_theme(True))
        theme_menu.addAction(light_action)
        theme_menu.addAction(dark_action)

        # Меню масштаба
        scale_menu = settings_menu.addMenu("Масштаб")
        zoom_in_action = QAction("Увеличить (+)", self)
        zoom_in_action.setShortcut(QKeySequence.ZoomIn)
        zoom_in_action.triggered.connect(self.zoom_in)
        zoom_out_action = QAction("Уменьшить (-)", self)
        zoom_out_action.setShortcut(QKeySequence.ZoomOut)
        zoom_out_action.triggered.connect(self.zoom_out)
        reset_zoom_action = QAction("Сбросить масштаб", self)
        reset_zoom_action.setShortcut("Ctrl+0")
        reset_zoom_action.triggered.connect(self.reset_zoom)
        scale_menu.addAction(zoom_in_action)
        scale_menu.addAction(zoom_out_action)
        scale_menu.addAction(reset_zoom_action)

        # Добавляем обработчики горячих клавиш для масштабирования
        zoom_in_shortcut = QShortcut(QKeySequence.ZoomIn, self)
        zoom_in_shortcut.activated.connect(self.zoom_in)
        zoom_out_shortcut = QShortcut(QKeySequence.ZoomOut, self)
        zoom_out_shortcut.activated.connect(self.zoom_out)
        reset_zoom_shortcut = QShortcut("Ctrl+0", self)
        reset_zoom_shortcut.activated.connect(self.reset_zoom)

    def set_theme(self, dark_mode: bool):
        """Устанавливает светлую или темную тему"""
        self.dark_mode = dark_mode
        self._apply_theme()

    def _apply_theme(self):
        """Применяет выбранную тему ко всему приложению"""
        app = QApplication.instance()

        if self.dark_mode:
            # Темная палитра
            palette = QPalette()
            palette.setColor(QPalette.Window, QColor(53, 53, 53))
            palette.setColor(QPalette.WindowText, Qt.white)
            palette.setColor(QPalette.Base, QColor(35, 35, 35))
            palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
            palette.setColor(QPalette.ToolTipBase, QColor(25, 25, 25))
            palette.setColor(QPalette.ToolTipText, Qt.white)
            palette.setColor(QPalette.Text, Qt.white)
            palette.setColor(QPalette.Button, QColor(53, 53, 53))
            palette.setColor(QPalette.ButtonText, Qt.white)
            palette.setColor(QPalette.BrightText, Qt.red)
            palette.setColor(QPalette.Highlight, QColor(142, 45, 197))
            palette.setColor(QPalette.HighlightedText, Qt.black)

            app.setPalette(palette)
            app.setStyleSheet("""
                QMainWindow, QDialog, QWidget {
                    background-color: #353535;
                }
                QTabWidget::pane {
                    border: 1px solid #444;
                    background: #353535;
                }
                QTabBar::tab {
                    background: #444;
                    color: white;
                    padding: 8px;
                    border: 1px solid #444;
                    border-bottom: none;
                    border-top-left-radius: 4px;
                    border-top-right-radius: 4px;
                }
                QTabBar::tab:selected {
                    background: #555;
                    border-color: #555;
                }
                QGroupBox {
                    border: 1px solid #555;
                    margin-top: 10px;
                    padding-top: 15px;
                    background: #353535;
                }
                QGroupBox::title {
                    color: white;
                    subcontrol-origin: margin;
                    left: 10px;
                }
                QLabel {
                    color: white;
                }
                QLineEdit, QTextEdit, QPlainTextEdit {
                    background: #454545;
                    color: white;
                    border: 1px solid #555;
                    padding: 5px;
                }
                QPushButton {
                    background: #555;
                    color: white;
                    border: 1px solid #555;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background: #666;
                }
                QListWidget, QTableWidget {
                    background: #454545;
                    color: white;
                    border: 1px solid #555;
                }
            """)
        else:
            # Светлая палитра
            app.setPalette(app.style().standardPalette())
            app.setStyleSheet("""
                QMainWindow, QDialog, QWidget {
                    background-color: #f0f0f0;
                }
                QTabWidget::pane {
                    border: 1px solid #C2C7CB;
                    background: #f0f0f0;
                }
                QTabBar::tab {
                    background: #F0F0F0;
                    color: black;
                    padding: 8px;
                    border: 1px solid #C2C7CB;
                    border-bottom: none;
                    border-top-left-radius: 4px;
                    border-top-right-radius: 4px;
                }
                QTabBar::tab:selected {
                    background: white;
                    border-color: #C2C7CB;
                }
                QGroupBox {
                    border: 1px solid #C2C7CB;
                    margin-top: 10px;
                    padding-top: 15px;
                    background: white;
                }
                QGroupBox::title {
                    color: #333;
                    subcontrol-origin: margin;
                    left: 10px;
                }
                QLabel {
                    color: black;
                }
                QLineEdit, QTextEdit, QPlainTextEdit {
                    background: white;
                    color: black;
                    border: 1px solid #C2C7CB;
                    padding: 5px;
                }
                QPushButton {
                    background: #f0f0f0;
                    color: black;
                    border: 1px solid #C2C7CB;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background: #e0e0e0;
                }
                QListWidget, QTableWidget {
                    background: white;
                    color: black;
                    border: 1px solid #C2C7CB;
                }
            """)

        # Обновляем стиль матриц
        self._update_matrices_style()

    def _update_percent_toggle_text(self):
        """Обновляет текст кнопки переключения режима"""
        self.percent_toggle.setText(
            "Показать абсолютные значения" if self.display_percent
            else "Показать в процентах"
        )

    def _toggle_percent_display(self):
        """Переключение между процентами и абсолютными значениями"""
        self.display_percent = not self.display_percent
        self._update_percent_toggle_text()
        self._display_results()

    def _update_matrices_style(self):
        """Обновляет стиль всех матриц при смене темы"""
        for key in self.matrix_entries:
            widget = self.matrix_entries[key]
            if isinstance(widget, (QLineEdit, QLabel)):
                if self.dark_mode:
                    widget.setStyleSheet(self._get_dark_matrix_cell_style())
                else:
                    widget.setStyleSheet(self._get_light_matrix_cell_style())

    def _get_dark_matrix_cell_style(self):
        """Возвращает стиль ячеек матрицы для темной темы"""
        return """
            QLineEdit, QLabel {
                border: 1px solid #555;
                padding: 5px;
                background-color: #454545;
                color: white;
                border-radius: 3px;
            }
            QLineEdit:focus {
                border: 2px solid #8E2DC5;
            }
        """

    def _get_light_matrix_cell_style(self):
        """Возвращает стиль ячеек матрицы для светлой темы"""
        return """
            QLineEdit, QLabel {
                border: 1px solid #ccc;
                padding: 5px;
                background-color: white;
                color: black;
                border-radius: 3px;
            }
            QLineEdit:focus {
                border: 2px solid #4CAF50;
            }
        """

    def zoom_in(self):
        """Увеличивает масштаб интерфейса"""
        new_scale = round(self.current_scale + self.SCALE_STEP, 1)
        if new_scale <= self.MAX_SCALE:
            self.current_scale = new_scale
            self._apply_scale()

    def zoom_out(self):
        """Уменьшает масштаб интерфейса"""
        new_scale = round(self.current_scale - self.SCALE_STEP, 1)
        if new_scale >= self.MIN_SCALE:
            self.current_scale = new_scale
            self._apply_scale()

    def reset_zoom(self):
        """Сбрасывает масштаб к значению по умолчанию"""
        self.current_scale = 1.0
        self._apply_scale()

    def _apply_scale(self):
        """Применяет текущий масштаб ко всему приложению"""
        app = QApplication.instance()

        # Устанавливаем шрифт для всего приложения
        font = app.font()
        base_size = 9 if self.dark_mode else 10
        font.setPointSize(int(base_size * self.current_scale))
        app.setFont(font)

        # Обновляем специальные виджеты
        self._update_special_widgets_fonts()

        # Обновляем все виджеты
        self.update()
        self.updateGeometry()

        # Принудительное обновление всех дочерних виджетов
        for widget in self.findChildren(QWidget):
            widget.updateGeometry()
            widget.update()

    def _update_special_widgets_fonts(self):
        """Обновляет шрифты специальных виджетов"""
        # Обновляем матрицы
        for key in self.matrix_entries:
            widget = self.matrix_entries[key]
            if isinstance(widget, (QLineEdit, QLabel)):
                font = widget.font()
                base_size = 8 if self.dark_mode else 9
                font.setPointSize(int(base_size * self.current_scale))
                widget.setFont(font)

        # Обновляем таблицы
        for table in self.findChildren(QTableWidget):
            font = table.font()
            font.setPointSize(int(10 * self.current_scale))
            table.setFont(font)
            table.resizeRowsToContents()
            table.resizeColumnsToContents()

    def wheelEvent(self, event):
        """Обработка масштабирования колесиком мыши с Ctrl"""
        if event.modifiers() & Qt.ControlModifier:
            angle = event.angleDelta().y()
            if angle > 0:
                self.zoom_in()
            else:
                self.zoom_out()
            event.accept()
        else:
            super().wheelEvent(event)

    def resizeEvent(self, event):
        """Обработчик изменения размера окна"""
        try:
            super().resizeEvent(event)
            if hasattr(self, 'centralWidget') and self.centralWidget():
                self.centralWidget().updateGeometry()
            QApplication.processEvents()
        except Exception as e:
            print(f"Ошибка в resizeEvent: {str(e)}")

    def force_maximize(self):
        """Принудительная максимизация окна"""
        try:
            if not self.isMaximized():
                self.showMaximized()
            # Дополнительная подгонка размеров
            self.resize(self.size().width(), self.size().height() - 1)
            self.resize(self.size().width() + 1, self.size().height())
        except Exception as e:
            print(f"Ошибка при максимизации: {str(e)}")

    def _create_widgets(self):
        """Создание всех виджетов интерфейса"""
        self.tabs = QTabWidget()

        # Создание вкладок
        self._create_hierarchy_tab()
        self._create_comparison_tab()
        self._create_consistency_tab()
        self._create_results_tab()

        # Блокировка вкладок
        for i in range(1, 4):
            self.tabs.setTabEnabled(i, False)

    def _create_hierarchy_tab(self):
        """Создание вкладки иерархии"""
        try:
            tab = QWidget()
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll_content = QWidget()
            scroll_layout = QVBoxLayout(scroll_content)

            # Группа выбора уровней иерархии
            levels_group = QGroupBox("1. Выберите количество уровней иерархии")
            levels_layout = QHBoxLayout()

            self.level_buttons = []
            btn_group = QButtonGroup()
            for i in range(1, 4):
                btn = QRadioButton(f"{i} {'уровня' if i > 1 else 'уровень'}")
                btn.setChecked(i == 3)
                btn.level = i
                btn.toggled.connect(self._update_hierarchy_levels)
                btn_group.addButton(btn)
                levels_layout.addWidget(btn)
                self.level_buttons.append(btn)

            levels_layout.addStretch()
            levels_group.setLayout(levels_layout)
            scroll_layout.addWidget(levels_group)

            # Группа альтернатив
            self.alt_group = QGroupBox("2. Ввод альтернатив")
            alt_layout = QHBoxLayout()

            self.alt_entry = QLineEdit()
            self.alt_entry.setFixedWidth(200)
            add_alt_btn = QPushButton("Добавить")
            add_alt_btn.clicked.connect(self._add_alternative)

            alt_layout.addWidget(QLabel("Альтернатива:"))
            alt_layout.addWidget(self.alt_entry)
            alt_layout.addWidget(add_alt_btn)
            alt_layout.addStretch()

            self.alt_list_widget = QWidget()
            self.alt_list_layout = QVBoxLayout(self.alt_list_widget)
            self.alt_list_layout.setAlignment(Qt.AlignTop)

            self.alt_group.setLayout(QVBoxLayout())
            self.alt_group.layout().addLayout(alt_layout)
            self.alt_group.layout().addWidget(self.alt_list_widget)

            # Группа критериев
            self.crit_group = QGroupBox("3. Ввод критериев")
            crit_layout = QHBoxLayout()

            self.crit_entry = QLineEdit()
            self.crit_entry.setFixedWidth(200)
            add_crit_btn = QPushButton("Добавить")
            add_crit_btn.clicked.connect(self._add_criterion)

            crit_layout.addWidget(QLabel("Критерий:"))
            crit_layout.addWidget(self.crit_entry)
            crit_layout.addWidget(add_crit_btn)
            crit_layout.addStretch()

            self.crit_list_widget = QWidget()
            self.crit_list_layout = QVBoxLayout(self.crit_list_widget)
            self.crit_list_layout.setAlignment(Qt.AlignTop)

            self.crit_group.setLayout(QVBoxLayout())
            self.crit_group.layout().addLayout(crit_layout)
            self.crit_group.layout().addWidget(self.crit_list_widget)

            # Группа типов критериев
            self.type_group = QGroupBox("4. Ввод видов критериев")
            type_layout = QVBoxLayout()

            # Поля ввода
            type_name_layout = QHBoxLayout()
            type_name_layout.addWidget(QLabel("Вид критериев:"))
            self.type_entry = QLineEdit()
            self.type_entry.setFixedWidth(150)
            self.type_entry.setPlaceholderText("Название вида")
            type_name_layout.addWidget(self.type_entry)
            type_name_layout.addStretch()
            type_layout.addLayout(type_name_layout)

            # Поиск критериев
            search_layout = QHBoxLayout()
            search_layout.addWidget(QLabel("Поиск критериев:"))
            self.criteria_search = QLineEdit()
            self.criteria_search.setPlaceholderText("Введите текст для поиска...")
            self.criteria_search.textChanged.connect(self._filter_criteria_list)
            search_layout.addWidget(self.criteria_search)
            type_layout.addLayout(search_layout)

            # Список критериев
            self.type_criteria = QListWidget()
            self.type_criteria.setSelectionMode(QListWidget.MultiSelection)
            self.type_criteria.setFixedHeight(100)
            type_layout.addWidget(self.type_criteria)

            # Кнопка добавления
            add_type_btn = QPushButton("Добавить вид критериев")
            add_type_btn.clicked.connect(self._add_criterion_type)
            type_layout.addWidget(add_type_btn)

            # Список добавленных видов
            self.type_list_widget = QWidget()
            self.type_list_layout = QVBoxLayout(self.type_list_widget)
            self.type_list_layout.setAlignment(Qt.AlignTop)
            type_layout.addWidget(self.type_list_widget)

            self.type_group.setLayout(type_layout)

            # Кнопка генерации
            gen_btn = QPushButton("Сгенерировать матрицы сравнений →")
            gen_btn.clicked.connect(self._generate_matrices)

            # Добавление виджетов в layout
            scroll_layout.addWidget(self.alt_group)
            scroll_layout.addWidget(self.crit_group)
            scroll_layout.addWidget(self.type_group)
            scroll_layout.addWidget(gen_btn)
            scroll_layout.addStretch()

            scroll.setWidget(scroll_content)
            tab_layout = QVBoxLayout(tab)
            tab_layout.addWidget(scroll)

            self.tabs.addTab(tab, "1. Определение элементов иерархии")
            return tab

        except Exception as e:
            QMessageBox.critical(None, "Ошибка", f"Ошибка создания вкладки иерархии: {str(e)}")
            raise

    def _update_hierarchy_levels(self):
        """Обновление видимости групп в зависимости от выбранных уровней"""
        try:
            self.selected_levels = self._get_selected_levels()

            if hasattr(self, 'alt_group'):
                self.alt_group.setVisible(self.selected_levels >= 1)
            if hasattr(self, 'crit_group'):
                self.crit_group.setVisible(self.selected_levels >= 2)
            if hasattr(self, 'type_group'):
                self.type_group.setVisible(self.selected_levels >= 3)
        except Exception as e:
            print(f"Ошибка в _update_hierarchy_levels: {str(e)}")

    def _get_selected_levels(self):
        """Получение количества выбранных уровней"""
        try:
            if hasattr(self, 'level_buttons'):
                for btn in self.level_buttons:
                    if btn.isChecked():
                        return btn.level
        except Exception as e:
            print(f"Ошибка в _get_selected_levels: {str(e)}")
        return 3  # Значение по умолчанию

    def _create_consistency_tab(self):
        """Создание вкладки проверки согласованности"""
        try:
            self.consistency_tab = QWidget()
            self.consistency_tab.setLayout(QVBoxLayout())
            self.consistency_tab.layout().setContentsMargins(10, 10, 10, 10)
            self.consistency_tab.layout().setSpacing(15)

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setFrameShape(QFrame.NoFrame)

            content_widget = QWidget()
            self.consistency_layout = QVBoxLayout(content_widget)
            self.consistency_layout.setContentsMargins(5, 5, 5, 5)
            self.consistency_layout.setSpacing(15)
            self.consistency_layout.setAlignment(Qt.AlignTop)

            scroll.setWidget(content_widget)
            self.consistency_tab.layout().addWidget(scroll)

            # Кнопка проверки согласованности
            self.recalc_btn = QPushButton("Проверить согласованность →")
            self.recalc_btn.setStyleSheet("""
                QPushButton {
                    font-weight: bold;
                    padding: 8px;
                    margin: 10px;
                    min-width: 200px;
                }
            """)
            self.recalc_btn.clicked.connect(self._check_all_consistency)

            btn_container = QWidget()
            btn_layout = QHBoxLayout(btn_container)
            btn_layout.addStretch()
            btn_layout.addWidget(self.recalc_btn)
            btn_layout.addStretch()

            self.consistency_tab.layout().addWidget(btn_container)
            self.tabs.addTab(self.consistency_tab, "3. Проверка согласованности")

        except Exception as e:
            QMessageBox.critical(None, "Ошибка", f"Ошибка создания вкладки согласованности: {str(e)}")
            raise

    def _create_comparison_tab(self):
        """Создание вкладки парных сравнений"""
        try:
            self.comp_tab = QWidget()
            self.comp_tab.setLayout(QVBoxLayout())

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll_content = QWidget()
            self.scroll_layout = QVBoxLayout(scroll_content)

            scroll.setWidget(scroll_content)
            self.comp_tab.layout().addWidget(scroll)

            self.tabs.addTab(self.comp_tab, "2. Метод парных сравнений")

        except Exception as e:
            QMessageBox.critical(None, "Ошибка", f"Ошибка создания вкладки сравнений: {str(e)}")
            raise

    def _check_all_consistency(self):
        """Проверка согласованности с устранением дублирования заголовков"""
        try:
            self._clear_layout(self.consistency_layout)
            all_consistent = True

            # 1. Проверка первого уровня (типы критериев) - только для 3 уровней
            if self.selected_levels >= 3 and 'criteria_types' in self.backend.matrices:
                title1 = QLabel("Первый уровень: Согласованность видов критериев")
                title1.setStyleSheet("font-weight: bold; font-size: 12pt;")
                self.consistency_layout.addWidget(title1)

                matrix = self.backend.matrices['criteria_types']
                consistency = self.backend.check_consistency(matrix)

                group = self._create_consistency_group("", consistency)  # Убрали дублирующий заголовок
                self.consistency_layout.addWidget(group)

                if "ТРЕБУЕТСЯ пересмотр" in consistency['status']:
                    all_consistent = False

            # 2. Проверка второго уровня (критерии) - для 2 и 3 уровней
            if self.selected_levels >= 2:
                title2 = QLabel("Второй уровень: Согласованность критериев")
                title2.setStyleSheet("font-weight: bold; font-size: 12pt; margin-top: 20px;")
                self.consistency_layout.addWidget(title2)

                if self.selected_levels >= 3:
                    # Для 3 уровней - проверка по типам критериев
                    for type_name in self.backend.criteria_types:
                        key = f'criteria_{type_name}'
                        if key in self.backend.matrices:
                            # Убрали подзаголовок "Вид критериев", оставили только название группы
                            group_title = f"Критерии ({type_name})"  # Измененная строка
                            matrix = self.backend.matrices[key]
                            consistency = self.backend.check_consistency(matrix)

                            group = self._create_consistency_group(group_title, consistency)
                            self.consistency_layout.addWidget(group)

                            if "ТРЕБУЕТСЯ пересмотр" in consistency['status']:
                                all_consistent = False
                else:
                    # Для 2 уровней - общая матрица критериев
                    if 'criteria' in self.backend.matrices:
                        matrix = self.backend.matrices['criteria']
                        consistency = self.backend.check_consistency(matrix)

                        group = self._create_consistency_group("Критерии", consistency)
                        self.consistency_layout.addWidget(group)

                        if "ТРЕБУЕТСЯ пересмотр" in consistency['status']:
                            all_consistent = False

            # 3. Проверка третьего уровня (альтернативы) - для всех уровней
            title3 = QLabel("Согласованность альтернатив")
            title3.setStyleSheet("font-weight: bold; font-size: 12pt; margin-top: 20px;")
            self.consistency_layout.addWidget(title3)

            if self.selected_levels >= 2:
                # Для 2 и 3 уровней - проверка по критериям
                for criterion in self.backend.criteria:
                    key = f'alternatives_{criterion}'
                    if key in self.backend.matrices:
                        # Убрали подзаголовок "Критерий", оставили только название группы
                        group_title = f"Альтернативы по критерию '{criterion}'"  # Измененная строка
                        matrix = self.backend.matrices[key]
                        consistency = self.backend.check_consistency(matrix)

                        group = self._create_consistency_group(group_title, consistency)
                        self.consistency_layout.addWidget(group)

                        if "ТРЕБУЕТСЯ пересмотр" in consistency['status']:
                            all_consistent = False
            else:
                # Для 1 уровня - общая матрица альтернатив
                if 'alternatives' in self.backend.matrices:
                    matrix = self.backend.matrices['alternatives']
                    consistency = self.backend.check_consistency(matrix)

                    group = self._create_consistency_group("Альтернативы", consistency)
                    self.consistency_layout.addWidget(group)

                    if "ТРЕБУЕТСЯ пересмотр" in consistency['status']:
                        all_consistent = False

            # Остальной код без изменений...
            self.tabs.setTabEnabled(3, all_consistent)

            if all_consistent:
                QMessageBox.information(self, "Проверка завершена",
                                        "Все матрицы имеют приемлемую согласованность!")
            else:
                QMessageBox.warning(self, "Внимание",
                                    "Некоторые матрицы требуют пересмотра!")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка проверки согласованности: {str(e)}")

    def _create_consistency_group(self, title: str, consistency: dict) -> QGroupBox:
        """Создает группу с информацией о согласованности"""
        group = QGroupBox(title)
        layout = QVBoxLayout()

        layout.addWidget(QLabel(f"λmax: {consistency['lambda_max']:.3f}"))
        layout.addWidget(QLabel(f"Индекс согласованности (ИС): {consistency['CI']:.3f}"))
        layout.addWidget(QLabel(f"Отношение согласованности (ОС): {consistency['CR']:.3f}"))

        status = QLabel(f"Статус: {consistency['status']}")
        if "ТРЕБУЕТСЯ пересмотр" in consistency['status']:
            status.setStyleSheet("color: red; font-weight: bold;")
        elif "Приемлемая согласованность" in consistency['status']:
            status.setStyleSheet("color: orange;")
        else:
            status.setStyleSheet("color: green;")

        layout.addWidget(status)
        layout.addWidget(QLabel("(ОС < 0.1 - отличная, ОС < 0.2 - приемлемая, ОС ≥ 0.2 - требует пересмотра)"))

        group.setLayout(layout)
        return group

    def _create_results_tab(self):
        """Создание вкладки результатов"""
        try:
            self.res_tab = QWidget()
            self.res_tab.setLayout(QVBoxLayout())

            self.res_controls_frame = QWidget()
            self.res_controls_layout = QHBoxLayout(self.res_controls_frame)

            self.res_display_frame = QScrollArea()
            self.res_display_frame.setWidgetResizable(True)
            self.res_display_content = QWidget()
            self.res_display_layout = QVBoxLayout(self.res_display_content)

            self.res_display_frame.setWidget(self.res_display_content)

            self.res_tab.layout().addWidget(self.res_controls_frame)
            self.res_tab.layout().addWidget(self.res_display_frame)

            self.tabs.addTab(self.res_tab, "4. Результаты анализа")

        except Exception as e:
            QMessageBox.critical(None, "Ошибка", f"Ошибка создания вкладки результатов: {str(e)}")
            raise

    def _add_alternative(self):
        """Добавление альтернативы"""
        try:
            if hasattr(self, 'alt_entry'):
                text = self.alt_entry.text().strip()
                if text and hasattr(self, 'backend') and self.backend.add_alternative(text):
                    self.alt_entry.clear()
                    self._update_alt_list()
                    self._update_criteria_listbox()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка добавления альтернативы: {str(e)}")

    def _add_criterion(self):
        """Добавление критерия"""
        try:
            if hasattr(self, 'crit_entry'):
                text = self.crit_entry.text().strip()
                if text and hasattr(self, 'backend') and self.backend.add_criterion(text):
                    self.crit_entry.clear()
                    self._update_crit_list()
                    self._update_criteria_listbox()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка добавления критерия: {str(e)}")

    def _add_criterion_type(self):
        """Добавление вида критериев"""
        try:
            if not hasattr(self, 'type_entry') or not hasattr(self, 'type_criteria'):
                return

            type_name = self.type_entry.text().strip()
            selected = [item.text() for item in self.type_criteria.selectedItems()]

            if not type_name:
                QMessageBox.warning(self, "Ошибка", "Введите название вида критериев")
                return
            if not selected:
                QMessageBox.warning(self, "Ошибка", "Выберите хотя бы один критерий")
                return

            if hasattr(self, 'backend') and self.backend.add_criterion_type(type_name, selected):
                self.type_entry.clear()
                if hasattr(self, 'criteria_search'):
                    self.criteria_search.clear()
                self.type_criteria.clearSelection()
                self._update_type_list()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка добавления вида критериев: {str(e)}")

    def _update_alt_list(self):
        """Обновление списка альтернатив"""
        try:
            if not hasattr(self, 'alt_list_layout') or not hasattr(self, 'backend'):
                return

            self._clear_layout(self.alt_list_layout)

            for i, alt in enumerate(getattr(self.backend, 'alternatives', [])):
                frame = QWidget()
                frame.setLayout(QHBoxLayout())

                label = QLabel(alt)
                btn = QPushButton("Удалить")
                btn.clicked.connect(lambda _, idx=i: self._remove_item('alternatives', idx))

                frame.layout().addWidget(label)
                frame.layout().addWidget(btn)
                frame.layout().addStretch()

                self.alt_list_layout.addWidget(frame)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка обновления списка альтернатив: {str(e)}")

    def _update_crit_list(self):
        """Обновление списка критериев"""
        try:
            if not hasattr(self, 'crit_list_layout') or not hasattr(self, 'backend'):
                return

            self._clear_layout(self.crit_list_layout)

            for i, crit in enumerate(getattr(self.backend, 'criteria', [])):
                frame = QWidget()
                frame.setLayout(QHBoxLayout())

                label = QLabel(crit)
                btn = QPushButton("Удалить")
                btn.clicked.connect(lambda _, idx=i: self._remove_item('criteria', idx))

                frame.layout().addWidget(label)
                frame.layout().addWidget(btn)
                frame.layout().addStretch()

                self.crit_list_layout.addWidget(frame)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка обновления списка критериев: {str(e)}")

    def _update_type_list(self):
        """Обновление списка видов критериев"""
        try:
            if not hasattr(self, 'type_list_layout') or not hasattr(self, 'backend'):
                return

            self._clear_layout(self.type_list_layout)

            for type_name, criteria in getattr(self.backend, 'criteria_types', {}).items():
                frame = QWidget()
                frame.setLayout(QHBoxLayout())

                label = QLabel(f"{type_name}: {', '.join(criteria)}")
                btn = QPushButton("Удалить")
                btn.clicked.connect(lambda _, tn=type_name: self._remove_criterion_type(tn))

                frame.layout().addWidget(label)
                frame.layout().addWidget(btn)
                frame.layout().addStretch()

                self.type_list_layout.addWidget(frame)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка обновления списка видов критериев: {str(e)}")

    def _filter_criteria_list(self):
        """Фильтрация списка критериев"""
        try:
            if not hasattr(self, 'criteria_search') or not hasattr(self, 'type_criteria'):
                return

            search_text = self.criteria_search.text().lower()
            for i in range(self.type_criteria.count()):
                item = self.type_criteria.item(i)
                if item:
                    item.setHidden(search_text not in item.text().lower())
        except Exception as e:
            print(f"Ошибка фильтрации списка критериев: {str(e)}")

    def _update_criteria_listbox(self):
        """Обновление списка критериев с сохранением фильтра"""
        try:
            if not hasattr(self, 'criteria_search') or not hasattr(self, 'type_criteria') or \
                    not hasattr(self, 'backend'):
                return

            current_search = self.criteria_search.text()
            self.type_criteria.clear()

            for crit in getattr(self.backend, 'criteria', []):
                self.type_criteria.addItem(crit)

            if current_search:
                self._filter_criteria_list()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка обновления списка критериев: {str(e)}")

    def _remove_item(self, item_type, index):
        """Удаление элемента (альтернативы или критерия)"""
        try:
            if not hasattr(self, 'backend'):
                return

            if item_type == 'alternatives':
                if 0 <= index < len(getattr(self.backend, 'alternatives', [])):
                    del self.backend.alternatives[index]
                    self._update_alt_list()
            elif item_type == 'criteria':
                if 0 <= index < len(getattr(self.backend, 'criteria', [])):
                    crit = self.backend.criteria[index]

                    # Удаление из типов критериев
                    for type_name in list(getattr(self.backend, 'criteria_types', {}).keys()):
                        if crit in self.backend.criteria_types[type_name]:
                            self.backend.criteria_types[type_name].remove(crit)
                            if not self.backend.criteria_types[type_name]:
                                del self.backend.criteria_types[type_name]

                    del self.backend.criteria[index]
                    self._update_crit_list()
                    self._update_type_list()
                    self._update_criteria_listbox()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка удаления элемента: {str(e)}")

    def _remove_criterion_type(self, type_name):
        """Удаление вида критериев"""
        try:
            if hasattr(self.backend, 'criteria_types') and type_name in self.backend.criteria_types:
                del self.backend.criteria_types[type_name]
                self._update_type_list()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка удаления вида критериев: {str(e)}")

    def _generate_matrices(self):
        """Генерация матриц сравнения с учетом выбранных уровней"""
        self.selected_levels = self._get_selected_levels()

        # Проверка обязательных полей
        if not self.backend.alternatives:
            QMessageBox.critical(self, "Ошибка", "Добавьте хотя бы одну альтернативу")
            return

        if self.selected_levels >= 2 and not self.backend.criteria:
            QMessageBox.critical(self, "Ошибка", "Добавьте хотя бы один критерий")
            return

        if self.selected_levels >= 3 and not self.backend.criteria_types:
            QMessageBox.critical(self, "Ошибка", "Добавьте хотя бы один вид критериев")
            return

        # Активация вкладки сравнений
        self.tabs.setTabEnabled(1, True)
        self.tabs.setCurrentIndex(1)
        self._setup_comparison_tab()

    def _setup_comparison_tab(self):
        """Настройка вкладки с матрицами сравнения с учетом уровней и количества матриц"""
        self._clear_layout(self.scroll_layout)

        main_frame = QFrame()
        main_layout = QVBoxLayout(main_frame)
        main_layout.setAlignment(Qt.AlignTop)  # Выравнивание по верхнему краю

        # Получаем количество матриц из backend
        results = self.backend.calculate_ahp(self.selected_levels)
        matrix_count = results.get('matrix_count', 1)

        # Стиль для заголовков
        title_style = """
            QLabel {
                font-weight: bold; 
                font-size: 12pt;
                margin-bottom: 10px;
            }
        """

        # Матрица сравнения альтернатив (для 1 уровня)
        if self.selected_levels == 1:
            title = QLabel("Матрица сравнения альтернатив")
            title.setStyleSheet(title_style)
            main_layout.addWidget(title, alignment=Qt.AlignTop)  # Выравнивание по верхнему краю

            self._create_matrix_ui(self.backend.alternatives, 'alternatives', main_layout)

        # Матрицы сравнения критериев и альтернатив (для 2 уровней)
        elif self.selected_levels == 2:
            # Матрица критериев
            title1 = QLabel("Матрица сравнения критериев")
            title1.setStyleSheet(title_style)
            main_layout.addWidget(title1, alignment=Qt.AlignTop)

            self._create_matrix_ui(self.backend.criteria, 'criteria', main_layout)

            # Матрицы альтернатив по каждому критерию
            title2 = QLabel("Матрицы сравнения альтернатив по критериям")
            title2.setStyleSheet(title_style)
            main_layout.addWidget(title2, alignment=Qt.AlignTop)

            for criterion in self.backend.criteria:
                subtitle = QLabel(f"Критерий: {criterion}")
                subtitle.setStyleSheet("font-weight: bold; margin-top: 15px;")
                main_layout.addWidget(subtitle, alignment=Qt.AlignTop)

                self._create_matrix_ui(
                    self.backend.alternatives,
                    f'alternatives_{criterion}',
                    main_layout
                )

        # Матрицы сравнения типов критериев, критериев и альтернатив (для 3 уровней)
        elif self.selected_levels >= 3:
            # Матрица типов критериев
            title1 = QLabel("Матрица сравнения видов критериев")
            title1.setStyleSheet(title_style)
            main_layout.addWidget(title1, alignment=Qt.AlignTop)

            type_names = list(self.backend.criteria_types.keys())
            self._create_matrix_ui(type_names, 'criteria_types', main_layout)

            # Матрицы критериев по типам
            title2 = QLabel("Матрицы сравнения критериев по видам")
            title2.setStyleSheet(title_style)
            main_layout.addWidget(title2, alignment=Qt.AlignTop)

            for type_name in self.backend.criteria_types:
                subtitle = QLabel(f"Вид критериев: {type_name}")
                subtitle.setStyleSheet("font-weight: bold; margin-top: 15px;")
                main_layout.addWidget(subtitle, alignment=Qt.AlignTop)

                criteria = self.backend.criteria_types[type_name]
                self._create_matrix_ui(criteria, f'criteria_{type_name}', main_layout)

            # Матрицы альтернатив по критериям
            title3 = QLabel("Матрицы сравнения альтернатив по критериям")
            title3.setStyleSheet(title_style)
            main_layout.addWidget(title3, alignment=Qt.AlignTop)

            for criterion in self.backend.criteria:
                subtitle = QLabel(f"Критерий: {criterion}")
                subtitle.setStyleSheet("font-weight: bold; margin-top: 15px;")
                main_layout.addWidget(subtitle, alignment=Qt.AlignTop)

                self._create_matrix_ui(
                    self.backend.alternatives,
                    f'alternatives_{criterion}',
                    main_layout
                )

        # Кнопка расчета (всегда внизу)
        calc_btn = QPushButton("Рассчитать приоритеты →")
        calc_btn.setStyleSheet("""
            QPushButton {
                font-weight: bold;
                padding: 10px;
                margin-top: 20px;
                min-width: 200px;
                background-color: #4CAF50;
                color: white;
                border: none;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        calc_btn.clicked.connect(self._calculate_priorities)

        btn_container = QWidget()
        btn_layout = QHBoxLayout(btn_container)
        btn_layout.addStretch()
        btn_layout.addWidget(calc_btn)
        btn_layout.addStretch()

        main_layout.addWidget(btn_container)
        self.scroll_layout.addWidget(main_frame)

    def _setup_results_controls(self):
        """Настройка элементов управления для результатов"""
        self._clear_layout(self.res_controls_layout)

        # Стиль для кнопок
        button_style = """
            QPushButton {
                padding: 8px 12px;
                border: 1px solid #ccc;
                border-radius: 4px;
                background: #f8f8f8;
                margin-right: 5px;
                min-width: 80px;
            }
            QPushButton:hover {
                background: #e8e8e8;
            }
            QPushButton:checked {
                background: #4CAF50;
                color: white;
                border-color: #3e8e41;
                font-weight: bold;
            }
        """

        # Группа кнопок отображения
        self.view_buttons = {}
        formats = [("График", "chart"), ("Таблица", "table"), ("Диаграмма", "diagram")]

        for text, mode in formats:
            btn = QPushButton(text)
            btn.setCheckable(True)
            btn.setChecked(self.result_display_mode == mode)
            btn.setStyleSheet(button_style)
            btn.clicked.connect(lambda _, m=mode: self._set_result_display_mode(m))
            self.res_controls_layout.addWidget(btn)
            self.view_buttons[mode] = btn

        # Кнопка переключения процентов/абсолютных значений
        self.percent_toggle = QPushButton()
        self._update_percent_toggle_text()
        self.percent_toggle.setCheckable(True)
        self.percent_toggle.setChecked(self.display_percent)
        self.percent_toggle.setStyleSheet(button_style)
        self.percent_toggle.clicked.connect(self._toggle_percent_display)
        self.res_controls_layout.addWidget(self.percent_toggle)

        # Кнопка экспорта результатов
        self.export_btn = QPushButton("Экспорт результатов")
        self.export_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 8px 12px;
                border-radius: 4px;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        self.export_btn.clicked.connect(self._export_results)
        self.res_controls_layout.addWidget(self.export_btn)

        self.res_controls_layout.addStretch()

    def _export_results(self):
        """Экспорт результатов анализа в файлы разных форматов"""
        try:
            if not hasattr(self, 'result_data') or not self.result_data:
                QMessageBox.warning(self, "Ошибка", "Нет данных для экспорта")
                return

            # Создаем диалог выбора файла
            file_dialog = QFileDialog()
            file_dialog.setAcceptMode(QFileDialog.AcceptSave)
            file_dialog.setWindowTitle("Экспорт результатов анализа")
            file_dialog.setNameFilters([
                "Документ Excel (*.xlsx)",
                "Документ Word (*.docx)",
                "Файл JSON (*.json)"
            ])

            if file_dialog.exec_() != QFileDialog.Accepted:
                return

            selected_filter = file_dialog.selectedNameFilter()
            file_path = file_dialog.selectedFiles()[0]

            # Добавляем расширение, если его нет
            if not file_path.endswith(('.xlsx', '.docx', '.json')):
                if "Excel" in selected_filter:
                    file_path += ".xlsx"
                elif "Word" in selected_filter:
                    file_path += ".docx"
                elif "JSON" in selected_filter:
                    file_path += ".json"

            # Экспорт в выбранный формат
            if "Excel" in selected_filter:
                self._export_to_excel(file_path)
            elif "Word" in selected_filter:
                self._export_to_word(file_path)
            elif "JSON" in selected_filter:
                self._export_to_json(file_path)

            QMessageBox.information(self, "Успешно", f"Результаты успешно экспортированы в файл:\n{file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте результатов: {str(e)}")

    def _export_to_excel(self, file_path):
        """Экспорт результатов в Excel с исправлением ошибки объединенных ячеек"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment
            from openpyxl.utils import get_column_letter

            wb = Workbook()
            ws = wb.active
            ws.title = "Результаты анализа"

            # Заголовок (без объединения ячеек)
            ws.append(["Результаты анализа методом МАИ", "", "", ""])
            ws['A1'].font = Font(bold=True, size=14)
            ws['A1'].alignment = Alignment(horizontal='center')

            # Приоритеты типов критериев (для 3 уровней)
            if self.selected_levels >= 3 and 'type_priority' in self.result_data['priorities']:
                ws.append(["Приоритеты видов критериев (Первый уровень)", "", "", ""])
                ws.append(["№", "Вид критериев", "Значение приоритета", "Процент"])

                for i, (type_name, value) in enumerate(zip(
                        self.backend.criteria_types.keys(),
                        self.result_data['priorities']['type_priority']
                ), 1):
                    ws.append([i, type_name, value, value * 100])

            # Приоритеты критериев (для 2 и 3 уровней)
            if self.selected_levels >= 2 and 'criteria_priority' in self.result_data['priorities']:
                ws.append([])
                ws.append(
                    ["Приоритеты критериев" + (" (Второй уровень)" if self.selected_levels >= 3 else ""), "", "", ""])
                ws.append(["№", "Критерий", "Значение приоритета", "Процент"])

                for i, (criterion, value) in enumerate(zip(
                        self.backend.criteria,
                        self.result_data['priorities']['criteria_priority']
                ), 1):
                    ws.append([i, criterion, value, value * 100])

            # Приоритеты альтернатив
            if 'alternatives_priority' in self.result_data['priorities']:
                ws.append([])
                ws.append(["Итоговые приоритеты альтернатив", "", "", ""])
                ws.append(["№", "Альтернатива", "Значение приоритета", "Процент"])

                for i, (alt, value) in enumerate(zip(
                        self.backend.alternatives,
                        self.result_data['priorities']['alternatives_priority']
                ), 1):
                    ws.append([i, alt, value, value * 100])

            # Форматирование столбцов (без работы с объединенными ячейками)
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter  # Получаем букву столбца
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2) * 1.2
                ws.column_dimensions[column].width = adjusted_width

            # Центрирование заголовков
            for row in ws.iter_rows(min_row=1, max_row=1):
                for cell in row:
                    cell.alignment = Alignment(horizontal='center')

            wb.save(file_path)

        except Exception as e:
            raise Exception(f"Ошибка экспорта в Excel: {str(e)}")

    def _export_to_word(self, file_path):
        """Экспорт результатов в Word"""
        try:
            doc = Document()

            # Заголовок
            title = doc.add_paragraph("Результаты анализа методом МАИ")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.size = Pt(14)
            title.runs[0].bold = True
            doc.add_paragraph()

            # Приоритеты типов критериев (для 3 уровней)
            if self.selected_levels >= 3 and 'type_priority' in self.result_data['priorities']:
                doc.add_paragraph("Приоритеты видов критериев (Первый уровень)", style='Heading 2')

                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '№'
                hdr_cells[1].text = 'Вид критериев'
                hdr_cells[2].text = 'Значение приоритета'
                hdr_cells[3].text = 'Процент'

                for i, (type_name, value) in enumerate(zip(
                        self.backend.criteria_types.keys(),
                        self.result_data['priorities']['type_priority']
                ), 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = type_name
                    row_cells[2].text = f"{value:.6f}"
                    row_cells[3].text = f"{value * 100:.2f}%"

                doc.add_paragraph()

            # Приоритеты критериев (для 2 и 3 уровней)
            if self.selected_levels >= 2 and 'criteria_priority' in self.result_data['priorities']:
                doc.add_paragraph("Приоритеты критериев" + (" (Второй уровень)" if self.selected_levels >= 3 else ""),
                                  style='Heading 2')

                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '№'
                hdr_cells[1].text = 'Критерий'
                hdr_cells[2].text = 'Значение приоритета'
                hdr_cells[3].text = 'Процент'

                for i, (criterion, value) in enumerate(zip(
                        self.backend.criteria,
                        self.result_data['priorities']['criteria_priority']
                ), 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = criterion
                    row_cells[2].text = f"{value:.6f}"
                    row_cells[3].text = f"{value * 100:.2f}%"

                doc.add_paragraph()

            # Приоритеты альтернатив
            if 'alternatives_priority' in self.result_data['priorities']:
                doc.add_paragraph("Итоговые приоритеты альтернатив", style='Heading 2')

                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '№'
                hdr_cells[1].text = 'Альтернатива'
                hdr_cells[2].text = 'Значение приоритета'
                hdr_cells[3].text = 'Процент'

                for i, (alt, value) in enumerate(zip(
                        self.backend.alternatives,
                        self.result_data['priorities']['alternatives_priority']
                ), 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = alt
                    row_cells[2].text = f"{value:.6f}"
                    row_cells[3].text = f"{value * 100:.2f}%"

            doc.save(file_path)

        except Exception as e:
            raise Exception(f"Ошибка экспорта в Word: {str(e)}")

    def _export_to_json(self, file_path):
        """Экспорт результатов в JSON с обработкой numpy массивов"""
        try:
            import json
            import numpy as np

            # Функция для преобразования numpy типов в стандартные Python типы
            def convert_numpy(obj):
                if isinstance(obj, np.ndarray):
                    return obj.tolist()
                elif isinstance(obj, (np.int_, np.intc, np.intp, np.int8,
                                      np.int16, np.int32, np.int64, np.uint8,
                                      np.uint16, np.uint32, np.uint64)):
                    return int(obj)
                elif isinstance(obj, (np.float_, np.float16, np.float32, np.float64)):
                    return float(obj)
                elif isinstance(obj, np.bool_):
                    return bool(obj)
                return obj

            # Подготовка данных для экспорта
            results = {
                'alternatives': self.backend.alternatives,
                'priorities': {},
                'consistency': {},
                'selected_levels': self.selected_levels
            }

            # Преобразуем приоритеты
            if hasattr(self.result_data, 'priorities'):
                for key, value in self.result_data['priorities'].items():
                    if isinstance(value, np.ndarray):
                        results['priorities'][key] = value.tolist()
                    else:
                        results['priorities'][key] = value

            # Преобразуем данные согласованности
            if hasattr(self.result_data, 'consistency'):
                for key, value in self.result_data['consistency'].items():
                    if isinstance(value, dict):
                        results['consistency'][key] = {
                            k: convert_numpy(v) for k, v in value.items()
                        }
                    else:
                        results['consistency'][key] = convert_numpy(value)

            if self.selected_levels >= 2:
                results['criteria'] = self.backend.criteria

            if self.selected_levels >= 3:
                results['criteria_types'] = list(self.backend.criteria_types.keys())

            # Сериализация с обработкой numpy объектов
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=4, default=convert_numpy)

        except Exception as e:
            raise Exception(f"Ошибка экспорта в JSON: {str(e)}")

    def _display_table_results(self):
        """Отображение результатов в виде таблицы с исправлениями"""
        try:
            # Очищаем предыдущие результаты
            self._clear_layout(self.res_display_layout)

            if not hasattr(self, 'result_data') or not self.result_data:
                QMessageBox.warning(self, "Нет данных", "Нет данных для отображения таблицы")
                return

            # Проверяем наличие необходимых данных
            if 'priorities' not in self.result_data:
                QMessageBox.warning(self, "Ошибка данных", "Отсутствуют данные о приоритетах")
                return

            # Создаем главный контейнер
            container = QWidget()
            layout = QVBoxLayout(container)

            # Приоритеты типов критериев (для 3 уровней)
            if self.selected_levels >= 3 and 'type_priority' in self.result_data['priorities']:
                type_names = list(self.backend.criteria_types.keys())
                if len(type_names) == len(self.result_data['priorities']['type_priority']):
                    self._create_priority_table(
                        container,
                        type_names,
                        self.result_data['priorities']['type_priority'],
                        "Приоритеты видов критериев (Первый уровень)"
                    )

            # Приоритеты критериев (для 2 и 3 уровней)
            if self.selected_levels >= 2 and 'criteria_priority' in self.result_data['priorities']:
                if len(self.backend.criteria) == len(self.result_data['priorities']['criteria_priority']):
                    title = "Приоритеты критериев" + (" (Второй уровень)" if self.selected_levels >= 3 else "")
                    self._create_priority_table(
                        container,
                        self.backend.criteria,
                        self.result_data['priorities']['criteria_priority'],
                        title
                    )

            # Приоритеты альтернатив
            if ('alternatives_priority' in self.result_data['priorities'] and
                    len(self.backend.alternatives) == len(self.result_data['priorities']['alternatives_priority'])):
                self._create_priority_table(
                    container,
                    self.backend.alternatives,
                    self.result_data['priorities']['alternatives_priority'],
                    "Итоговые приоритеты альтернатив",
                    show_percent=True
                )

            # Если таблицы не были созданы
            if layout.count() == 0:
                QMessageBox.information(self, "Информация", "Нет данных для отображения таблицы")
            else:
                # Добавляем контейнер с таблицами в основной layout
                scroll = QScrollArea()
                scroll.setWidgetResizable(True)
                scroll.setWidget(container)
                self.res_display_layout.addWidget(scroll)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при отображении таблицы: {str(e)}")

    def _create_priority_table(self, parent_layout, labels, values, title, show_percent=False):
        """Создает таблицу с приоритетами с исправленной обработкой значений"""
        try:
            # Преобразуем значения в numpy array и проверяем их
            values = np.array(values, dtype=float)
            if values.size == 0:
                raise ValueError("Массив значений пуст")

            group = QGroupBox(title)
            layout = QVBoxLayout(group)

            table = QTableWidget()
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels(["№", "Элемент", "Значение", "Процент"])
            table.verticalHeader().setVisible(False)
            table.setEditTriggers(QTableWidget.NoEditTriggers)

            table.setRowCount(len(labels))

            # Находим максимальное значение безопасным способом
            max_value = np.max(values) if values.size > 0 else 0

            for row, (label, value) in enumerate(zip(labels, values)):
                # Номер
                item_num = QTableWidgetItem(str(row + 1))
                item_num.setTextAlignment(Qt.AlignCenter)

                # Название
                item_label = QTableWidgetItem(label)
                item_label.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)

                # Значение
                item_value = QTableWidgetItem(f"{value:.6f}")
                item_value.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

                # Процент
                percent = value * 100
                item_percent = QTableWidgetItem(f"{percent:.2f}%")
                item_percent.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

                # Подсветка максимального значения (с безопасным сравнением)
                if not np.isnan(value) and np.allclose(value, max_value):
                    for item in [item_num, item_label, item_value, item_percent]:
                        item.setBackground(QColor(230, 255, 230))
                        item.setForeground(QColor(0, 100, 0))
                        font = item.font()
                        font.setBold(True)
                        item.setFont(font)

                table.setItem(row, 0, item_num)
                table.setItem(row, 1, item_label)
                table.setItem(row, 2, item_value)
                table.setItem(row, 3, item_percent)

            table.resizeColumnsToContents()
            table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
            layout.addWidget(table)
            parent_layout.addWidget(group)

        except Exception as e:
            error_msg = f"Ошибка создания таблицы: {str(e)}"
            print(error_msg)  # Логируем ошибку
            raise Exception(error_msg)

    def _create_priority_table(self, parent, labels, values, title, show_percent=False):
        """Создание таблицы с приоритетами"""
        try:
            # Преобразуем значения в numpy array
            values = np.array(values, dtype=float)

            # Создаем группу для таблицы
            group = QGroupBox(title)
            layout = QVBoxLayout(group)

            # Создаем таблицу
            table = QTableWidget()
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels(
                ["№", "Элемент", "Значение приоритета", "Процент" if show_percent else "Доля"])
            table.verticalHeader().setVisible(False)
            table.setEditTriggers(QTableWidget.NoEditTriggers)
            table.setSelectionMode(QTableWidget.SingleSelection)

            # Устанавливаем количество строк
            table.setRowCount(len(labels))

            # Находим максимальное значение для подсветки
            max_value = np.max(values) if values.size > 0 else 0

            # Заполняем таблицу данными
            for row, (label, value) in enumerate(zip(labels, values)):
                # Номер
                item_num = QTableWidgetItem(str(row + 1))
                item_num.setTextAlignment(Qt.AlignCenter)

                # Название элемента
                item_label = QTableWidgetItem(label)
                item_label.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)

                # Значение приоритета
                item_value = QTableWidgetItem(f"{value:.6f}")
                item_value.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

                # Процентное значение
                percent_value = value * 100
                item_percent = QTableWidgetItem(f"{percent_value:.2f}%")
                item_percent.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

                # Подсветка максимального значения
                if not np.isnan(value) and np.allclose(value, max_value):
                    for item in [item_num, item_label, item_value, item_percent]:
                        item.setBackground(QColor(230, 255, 230))
                        item.setForeground(QColor(0, 100, 0))
                        font = item.font()
                        font.setBold(True)
                        item.setFont(font)

                # Добавляем ячейки в таблицу
                table.setItem(row, 0, item_num)
                table.setItem(row, 1, item_label)
                table.setItem(row, 2, item_value)
                table.setItem(row, 3, item_percent)

            # Настраиваем размеры столбцов
            table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
            table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)

            # Устанавливаем минимальную высоту таблицы
            table.setMinimumHeight(min(300, len(labels) * 30 + 50))

            # Добавляем таблицу в группу
            layout.addWidget(table)

            # Добавляем группу в родительский layout
            parent.layout().addWidget(group)

        except Exception as e:
            raise Exception(f"Ошибка создания таблицы приоритетов: {str(e)}")

    def _toggle_percent_display(self):
        """Переключение между процентами и абсолютными значениями"""
        try:
            self.display_percent = not self.display_percent
            if hasattr(self, 'percent_toggle'):
                self.percent_toggle.setChecked(self.display_percent)
                self.percent_toggle.setText(
                    "Показать в процентах" if not self.display_percent
                    else "Показать абсолютные значения"
                )
            self._display_results()
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Ошибка переключения режима: {str(e)}")

    def _create_matrix_ui(self, items, matrix_key, parent_layout):
        """Создание интерфейса матрицы сравнения с улучшенным стилем"""
        try:
            if not items or not isinstance(items, list):
                raise ValueError("Некорректный список элементов для матрицы")

            frame = QFrame()
            frame.setFrameShape(QFrame.StyledPanel)
            frame.setStyleSheet("""
                QFrame {
                    background-color: %s;
                    border-radius: 5px;
                    padding: 10px;
                    margin-bottom: 15px;
                }
            """ % ("#353535" if self.dark_mode else "#f9f9f9"))

            layout = QVBoxLayout(frame)
            layout.setContentsMargins(10, 10, 10, 10)
            layout.setSpacing(10)

            # Подсказка по шкале Саати
            saaty_tip = QLabel(
                "Шкала Саати:\n"
                "1 — равная важность; 3 — умеренное превосходство;\n"
                "5 — существенное превосходство; 7 — значительное превосходство;\n"
                "9 — абсолютное превосходство; 2,4,6,8 — промежуточные значения"
            )
            saaty_tip.setStyleSheet("""
                QLabel {
                    background-color: %s;
                    border: 1px solid %s;
                    padding: 8px;
                    border-radius: 4px;
                    font-size: %dpt;
                    margin-bottom: 10px;
                    color: %s;
                }
            """ % (
                "#444" if self.dark_mode else "#f0f0f0",
                "#555" if self.dark_mode else "#ddd",
                int(10 * self.current_scale),
                "white" if self.dark_mode else "black"
            ))
            layout.addWidget(saaty_tip)

            grid = QGridLayout()
            grid.setSpacing(5)
            grid.setContentsMargins(5, 5, 5, 5)

            n = len(items)
            cell_size = int((90 if n <= 5 else 70) * self.current_scale)

            # Устанавливаем стили для ячеек и заголовков
            header_style = self._get_matrix_header_style()
            cell_style = self._get_dark_matrix_cell_style() if self.dark_mode else self._get_light_matrix_cell_style()

            # Заголовки столбцов
            for j in range(n):
                label = QLabel(items[j] if j < len(items) else "")
                label.setStyleSheet(header_style)
                grid.addWidget(label, 0, j + 1)

            # Заполнение матрицы
            for i in range(n):
                # Заголовок строки
                row_label = QLabel(items[i] if i < len(items) else "")
                row_label.setStyleSheet(header_style)
                grid.addWidget(row_label, i + 1, 0)

                for j in range(n):
                    if i == j:
                        label = QLabel("1")
                        label.setStyleSheet(cell_style)
                        grid.addWidget(label, i + 1, j + 1)
                    elif i < j:
                        entry = QLineEdit()
                        entry.setStyleSheet(cell_style)
                        validator = QRegExpValidator(QRegExp(r"^([1-9]|1/[1-9])$"))
                        entry.setValidator(validator)
                        entry.setToolTip("Введите значение по шкале Саати (1-9 или 1/1-1/9)")

                        def make_lambda(key, row, col):
                            return lambda: self._safe_update_reciprocal(key, row, col)

                        entry.editingFinished.connect(make_lambda(matrix_key, i, j))
                        grid.addWidget(entry, i + 1, j + 1)
                        self.matrix_entries[(matrix_key, i, j)] = entry
                    else:
                        label = QLabel("")
                        label.setStyleSheet(cell_style)
                        grid.addWidget(label, i + 1, j + 1)
                        self.matrix_entries[(matrix_key, i, j)] = label

            layout.addLayout(grid)
            parent_layout.addWidget(frame)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка создания матрицы: {str(e)}")
            raise

    def _get_matrix_header_style(self):
        """Возвращает стиль заголовков матрицы"""
        cell_size = 90 if len(self.backend.alternatives) <= 5 else 70
        if self.dark_mode:
            return f"""
                QLabel {{
                    font-weight: bold;
                    border: 1px solid #666;
                    padding: 5px;
                    background-color: #444;
                    color: white;
                    min-width: {cell_size}px;
                    min-height: 30px;
                    text-align: center;
                    border-radius: 3px;
                }}
            """
        else:
            return f"""
                QLabel {{
                    font-weight: bold;
                    border: 1px solid #999;
                    padding: 5px;
                    background-color: #e0e0e0;
                    color: black;
                    min-width: {cell_size}px;
                    min-height: 30px;
                    text-align: center;
                    border-radius: 3px;
                }}
            """

    def _clear_layout(self, layout):
        """Безопасная очистка layout"""
        try:
            if layout is None:
                return

            while layout.count():
                item = layout.takeAt(0)
                widget = item.widget()
                if widget:
                    widget.deleteLater()
                else:
                    sublayout = item.layout()
                    if sublayout:
                        self._clear_layout(sublayout)
        except Exception as e:
            print(f"Ошибка очистки layout: {str(e)}")

    def _safe_update_reciprocal(self, matrix_key, i, j):
        """Безопасное обновление обратного значения"""
        try:
            if not hasattr(self, 'matrix_entries'):
                return

            entry = self.matrix_entries.get((matrix_key, i, j))
            recip_entry = self.matrix_entries.get((matrix_key, j, i))

            if not entry or not recip_entry:
                return

            value = entry.text().strip()
            if not value:
                return

            if value.startswith("1/"):
                recip_value = value[2:]
            else:
                recip_value = f"1/{value}" if value != "1" else "1"

            if isinstance(recip_entry, QLabel):
                recip_entry.setText(recip_value)
        except Exception as e:
            print(f"Ошибка обновления обратного значения: {str(e)}")

    def _calculate_priorities(self):
        """Расчет приоритетов с улучшенной обработкой ошибок"""
        try:
            # Проверка заполнения матриц
            if not self._check_all_matrices_filled():
                QMessageBox.warning(self, "Ошибка",
                                    "Не все матрицы сравнений заполнены!\n"
                                    "Заполните все необходимые матрицы перед расчетом.")
                return

            # Сбор данных для расчета
            matrices = {}
            errors = []

            # Для 3 уровней - матрица типов критериев
            if self.selected_levels >= 3:
                type_names = list(self.backend.criteria_types.keys())
                if not type_names:
                    errors.append("Не заданы типы критериев")
                else:
                    type_comparisons = self._collect_comparisons('criteria_types', type_names)
                    if type_comparisons:
                        matrix = self.backend.build_matrix(type_names, type_comparisons)
                        if matrix is not None:
                            matrices['criteria_types'] = matrix
                        else:
                            errors.append("Ошибка в матрице сравнения видов критериев")

            # Для 2 и 3 уровней - матрицы критериев
            if self.selected_levels >= 2 and not errors:
                if self.selected_levels >= 3:
                    # Для 3 уровней - матрицы критериев по типам
                    for type_name, type_criteria in self.backend.criteria_types.items():
                        comparisons = self._collect_comparisons(f'criteria_{type_name}', type_criteria)
                        if comparisons:
                            matrix = self.backend.build_matrix(type_criteria, comparisons)
                            if matrix is not None:
                                matrices[f'criteria_{type_name}'] = matrix
                            else:
                                errors.append(f"Ошибка в матрице критериев для вида '{type_name}'")
                else:
                    # Для 2 уровней - одна матрица критериев
                    criteria = self.backend.criteria
                    if not criteria:
                        errors.append("Не заданы критерии")
                    else:
                        comparisons = self._collect_comparisons('criteria', criteria)
                        if comparisons:
                            matrix = self.backend.build_matrix(criteria, comparisons)
                            if matrix is not None:
                                matrices['criteria'] = matrix
                            else:
                                errors.append("Ошибка в матрице сравнения критериев")

            # Для всех уровней - матрицы альтернатив
            if self.backend.alternatives and not errors:
                if self.selected_levels >= 2:
                    # Для 2 и 3 уровней - матрицы альтернатив по критериям
                    for criterion in self.backend.criteria:
                        comparisons = self._collect_comparisons(f'alternatives_{criterion}', self.backend.alternatives)
                        if comparisons:
                            matrix = self.backend.build_matrix(self.backend.alternatives, comparisons)
                            if matrix is not None:
                                matrices[f'alternatives_{criterion}'] = matrix
                            else:
                                errors.append(f"Ошибка в матрице альтернатив для критерия '{criterion}'")
                else:
                    # Для 1 уровня - одна матрица альтернатив
                    comparisons = self._collect_comparisons('alternatives', self.backend.alternatives)
                    if comparisons:
                        matrix = self.backend.build_matrix(self.backend.alternatives, comparisons)
                        if matrix is not None:
                            matrices['alternatives'] = matrix
                        else:
                            errors.append("Ошибка в матрице сравнения альтернатив")

            # Проверка ошибок перед расчетом
            if errors:
                QMessageBox.critical(self, "Ошибка", "\n".join(errors))
                return

            # Сохранение матриц и расчет AHP
            self.backend.matrices = matrices
            results = self.backend.calculate_ahp(self.selected_levels)

            if results is None:
                QMessageBox.critical(self, "Ошибка", "Не удалось рассчитать приоритеты")
                return

            if 'errors' in results and results['errors']:
                QMessageBox.warning(self, "Предупреждение", "\n".join(results['errors']))

            self._show_results(results)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при расчетах:\n{str(e)}")

    def _collect_comparisons(self, matrix_key, items):
        """Собирает сравнения для матрицы из полей ввода"""
        comparisons = {}
        n = len(items)
        for i in range(n):
            for j in range(i + 1, n):
                entry = self.matrix_entries.get((matrix_key, i, j))
                if entry and isinstance(entry, QLineEdit):
                    value = entry.text().strip()
                    if value:
                        comparisons[(i, j)] = value
        return comparisons


    def _check_all_matrices_filled(self):
        """Проверка заполнения всех матриц"""
        try:
            # Проверка матрицы типов критериев (для 3 уровней)
            if self.selected_levels >= 3:
                type_names = list(getattr(self.backend, 'criteria_types', {}).keys())
                for i in range(len(type_names)):
                    for j in range(i + 1, len(type_names)):
                        entry = self.matrix_entries.get(('criteria_types', i, j))
                        if not entry or not entry.text().strip():
                            return False

            # Проверка матриц критериев (для 2 и 3 уровней)
            if self.selected_levels >= 2:
                if self.selected_levels >= 3:
                    # Для 3 уровней - проверка матриц критериев по типам
                    for type_name in getattr(self.backend, 'criteria_types', {}):
                        criteria = self.backend.criteria_types[type_name]
                        for i in range(len(criteria)):
                            for j in range(i + 1, len(criteria)):
                                entry = self.matrix_entries.get((f'criteria_{type_name}', i, j))
                                if not entry or not entry.text().strip():
                                    return False
                else:
                    # Для 2 уровней - проверка одной матрицы критериев
                    criteria = getattr(self.backend, 'criteria', [])
                    for i in range(len(criteria)):
                        for j in range(i + 1, len(criteria)):
                            entry = self.matrix_entries.get(('criteria', i, j))
                            if not entry or not entry.text().strip():
                                return False

            # Проверка матриц альтернатив (для всех уровней)
            if getattr(self.backend, 'alternatives', []):
                if self.selected_levels >= 2:
                    # Для 2 и 3 уровней - проверка матриц альтернатив по критериям
                    for criterion in getattr(self.backend, 'criteria', []):
                        for i in range(len(self.backend.alternatives)):
                            for j in range(i + 1, len(self.backend.alternatives)):
                                entry = self.matrix_entries.get((f'alternatives_{criterion}', i, j))
                                if not entry or not entry.text().strip():
                                    return False
                else:
                    # Для 1 уровня - проверка одной матрицы альтернатив
                    for i in range(len(self.backend.alternatives)):
                        for j in range(i + 1, len(self.backend.alternatives)):
                            entry = self.matrix_entries.get(('alternatives', i, j))
                            if not entry or not entry.text().strip():
                                return False

            return True
        except Exception as e:
            print(f"Ошибка проверки матриц: {str(e)}")
            return False


    def _show_results(self, results):
        """Отображение результатов с проверкой согласованности"""
        try:
            # Сохраняем результаты
            self.result_data = results

            # Проверка согласованности
            all_consistent = True
            if 'consistency' in results:
                for data in results['consistency'].values():
                    if data['status'] not in ["Отличная согласованность", "Приемлемая согласованность"]:
                        all_consistent = False
                        break

            # Управление вкладками
            self.tabs.setTabEnabled(2, True)  # Вкладка согласованности
            self.tabs.setTabEnabled(3, all_consistent)  # Вкладка результатов
            self.tabs.setCurrentIndex(2 if not all_consistent else 3)

            # Настройка отображения
            if all_consistent:
                self._setup_results_controls()
                self._display_results()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка отображения результатов: {str(e)}")


    def _set_result_display_mode(self, mode):
        """Установка режима отображения результатов"""
        try:
            self.result_display_mode = mode

            # Обновляем состояние кнопок
            if hasattr(self, 'view_buttons'):
                for btn in self.view_buttons.values():
                    btn.setChecked(False)
                if mode in self.view_buttons:
                    self.view_buttons[mode].setChecked(True)

            # Обновляем отображение
            self._display_results()

        except Exception as e:
            print(f"Ошибка при смене режима отображения: {str(e)}")


    def _display_results(self):
        """Отображение результатов в выбранном режиме"""
        try:
            self._clear_layout(self.res_display_layout)

            if not hasattr(self, 'result_data') or not self.result_data:
                QMessageBox.warning(self, "Нет данных", "Нет данных для отображения")
                return

            # Создаем контейнер с прокруткой
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            content = QWidget()
            layout = QVBoxLayout(content)
            layout.setAlignment(Qt.AlignTop)

            # Выбираем способ отображения
            if self.result_display_mode == "table":
                self._display_table_results(layout)
            elif self.result_display_mode == "chart":
                self._display_chart_results(layout)
            elif self.result_display_mode == "diagram":
                self._display_diagram_results(layout)

            scroll.setWidget(content)
            self.res_display_layout.addWidget(scroll)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка отображения: {str(e)}")


    def _display_all_tables(self, layout):
        """Отображение всех данных в табличном виде"""
        # Типы критериев (для 3 уровней)
        if self.selected_levels >= 3 and 'type_priority' in self.result_data['priorities']:
            types = list(self.backend.criteria_types.keys())
            if len(types) == len(self.result_data['priorities']['type_priority']):
                self._create_table(
                    types,
                    self.result_data['priorities']['type_priority'],
                    "Приоритеты видов критериев (Первый уровень)"
                )

        # Критерии (для 2 и 3 уровней)
        if 'criteria_priority' in self.result_data['priorities']:
            criteria = self.backend.criteria
            if len(criteria) == len(self.result_data['priorities']['criteria_priority']):
                title = "Приоритеты критериев" + (" (Второй уровень)" if self.selected_levels >= 3 else "")
                self._create_table(
                    criteria,
                    self.result_data['priorities']['criteria_priority'],
                    title
                )

        # Альтернативы
        if 'alternatives_priority' in self.result_data['priorities']:
            alts = self.backend.alternatives
            if len(alts) == len(self.result_data['priorities']['alternatives_priority']):
                self._create_table(
                    alts,
                    self.result_data['priorities']['alternatives_priority'],
                    "Итоговые приоритеты альтернатив",
                    show_percent=True
                )


    def _display_all_charts(self, layout):
        """Отображение всех данных в виде столбчатых графиков с прокруткой"""
        try:
            # Очищаем предыдущие графики
            self._clear_layout(layout)

            # Создаем контейнер с вертикальным layout
            container = QWidget()
            container_layout = QVBoxLayout(container)
            container_layout.setSpacing(30)  # Отступ между графиками
            container_layout.setContentsMargins(20, 20, 20, 20)  # Отступы по краям

            # Счетчик добавленных графиков
            graphs_added = 0

            # 1. Типы критериев (для 3+ уровней)
            if self.selected_levels >= 3 and 'type_priority' in self.result_data['priorities']:
                types = list(self.backend.criteria_types.keys())
                values = self.result_data['priorities']['type_priority']

                if len(types) == len(values):
                    fig = plt.figure(figsize=(12, 6), facecolor='#f8f8f8')
                    ax = fig.add_subplot(111)

                    bars = ax.bar(types, values, color='#4C72B0', alpha=0.8)
                    ax.set_title("ПРИОРИТЕТЫ ВИДОВ КРИТЕРИЕВ",
                                 fontsize=16, pad=20, fontweight='bold')
                    ax.set_ylabel("Значение приоритета", fontsize=14)
                    ax.grid(axis='y', linestyle='--', alpha=0.5)
                    ax.set_ylim(0, max(values) * 1.15 if len(values) > 0 else 1)

                    # Поворот подписей и добавление значений
                    plt.xticks(rotation=45, ha='right')
                    for bar in bars:
                        height = bar.get_height()
                        ax.text(bar.get_x() + bar.get_width() / 2., height + 0.01,
                                f"{height:.3f}", ha='center', va='bottom', fontsize=12)

                    canvas = FigureCanvas(fig)
                    canvas.setMinimumHeight(500)  # Фиксированная высота
                    container_layout.addWidget(canvas)
                    graphs_added += 1

            # 2. Критерии (для 2 и 3 уровней)
            if 'criteria_priority' in self.result_data['priorities']:
                criteria = self.backend.criteria
                values = self.result_data['priorities']['criteria_priority']

                if len(criteria) == len(values):
                    fig = plt.figure(figsize=(12, 6), facecolor='#f8f8f8')
                    ax = fig.add_subplot(111)

                    bars = ax.bar(criteria, values, color='#55A868', alpha=0.8)
                    title = "ПРИОРИТЕТЫ КРИТЕРИЕВ" + (" (Второй уровень)" if self.selected_levels >= 3 else "")
                    ax.set_title(title, fontsize=16, pad=20, fontweight='bold')
                    ax.set_ylabel("Значение приоритета", fontsize=14)
                    ax.grid(axis='y', linestyle='--', alpha=0.5)
                    ax.set_ylim(0, max(values) * 1.15 if len(values) > 0 else 1)

                    # Поворот подписей и добавление значений
                    plt.xticks(rotation=45, ha='right')
                    for bar in bars:
                        height = bar.get_height()
                        ax.text(bar.get_x() + bar.get_width() / 2., height + 0.01,
                                f"{height:.3f}", ha='center', va='bottom', fontsize=12)

                    canvas = FigureCanvas(fig)
                    canvas.setMinimumHeight(500)
                    container_layout.addWidget(canvas)
                    graphs_added += 1

            # 3. Альтернативы
            if 'alternatives_priority' in self.result_data['priorities']:
                alts = self.backend.alternatives
                values = self.result_data['priorities']['alternatives_priority']

                if len(alts) == len(values):
                    fig = plt.figure(figsize=(12, 6), facecolor='#f8f8f8')
                    ax = fig.add_subplot(111)

                    bars = ax.bar(alts, values, color='#C44E52', alpha=0.8)
                    ax.set_title("ПРИОРИТЕТЫ АЛЬТЕРНАТИВ",
                                 fontsize=16, pad=20, fontweight='bold')
                    ax.set_ylabel("Значение приоритета", fontsize=14)
                    ax.grid(axis='y', linestyle='--', alpha=0.5)
                    ax.set_ylim(0, max(values) * 1.15 if len(values) > 0 else 1)

                    # Поворот подписей и добавление значений
                    plt.xticks(rotation=45, ha='right')
                    for bar in bars:
                        height = bar.get_height()
                        ax.text(bar.get_x() + bar.get_width() / 2., height + 0.01,
                                f"{height:.3f}", ha='center', va='bottom', fontsize=12)

                    canvas = FigureCanvas(fig)
                    canvas.setMinimumHeight(500)
                    container_layout.addWidget(canvas)
                    graphs_added += 1

            # Добавляем контейнер в ScrollArea только если есть графики
            if graphs_added > 0:
                scroll = QScrollArea()
                scroll.setWidgetResizable(True)
                scroll.setWidget(container)
                layout.addWidget(scroll)
            else:
                QMessageBox.information(self, "Информация", "Нет данных для отображения графиков")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка создания графиков: {str(e)}")


    def _display_chart_results(self, layout):
        """Отображение результатов в виде столбчатых диаграмм с учетом уровней иерархии"""
        try:
            self._clear_layout(layout)

            container = QWidget()
            container_layout = QVBoxLayout(container)
            container_layout.setSpacing(30)
            container_layout.setContentsMargins(20, 20, 20, 20)

            # Возвращаем оригинальный стиль
            plt.style.use('default')

            def create_bar_chart(labels, values, title, color):
                """Создает одну столбчатую диаграмму"""
                fig = plt.figure(figsize=(12, 6))
                ax = fig.add_subplot(111)

                if self.display_percent:
                    values = [v * 100 for v in values]
                    ylabel = "Приоритет, %"
                    fmt = lambda v: f"{v:.1f}%"
                else:
                    ylabel = "Значение приоритета"
                    fmt = lambda v: f"{v:.3f}"

                bars = ax.bar(labels, values, color=color, alpha=0.8)
                ax.set_title(title, fontsize=16, pad=20, fontweight='bold')
                ax.set_ylabel(ylabel, fontsize=14)
                ax.grid(axis='y', linestyle='--', alpha=0.5)

                if len(values) > 0:
                    ax.set_ylim(0, max(values) * 1.15)

                plt.xticks(rotation=45, ha='right')

                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.01,
                            fmt(height),
                            ha='center', va='bottom', fontsize=12)

                canvas = FigureCanvas(fig)
                canvas.setMinimumHeight(500)
                container_layout.addWidget(canvas)

            # Для 3 уровня - график типов критериев (первый уровень)
            if self.selected_levels >= 3 and 'type_priority' in self.result_data['priorities']:
                types = list(self.backend.criteria_types.keys())
                values = self.result_data['priorities']['type_priority']
                if len(types) == len(values):
                    create_bar_chart(types, values,
                                     "ПРИОРИТЕТЫ ВИДОВ КРИТЕРИЕВ (Первый уровень)",
                                     '#4C72B0')

            # Для 2 и 3 уровней - график критериев
            if self.selected_levels >= 2 and 'criteria_priority' in self.result_data['priorities']:
                criteria = self.backend.criteria
                values = self.result_data['priorities']['criteria_priority']
                if len(criteria) == len(values):
                    title = "ПРИОРИТЕТЫ КРИТЕРИЕВ" + (" (Второй уровень)" if self.selected_levels >= 3 else "")
                    create_bar_chart(criteria, values, title, '#55A868')

            # Для всех уровней - график альтернатив
            if 'alternatives_priority' in self.result_data['priorities']:
                alts = self.backend.alternatives
                values = self.result_data['priorities']['alternatives_priority']
                if len(alts) == len(values):
                    create_bar_chart(alts, values,
                                     "ПРИОРИТЕТЫ АЛЬТЕРНАТИВ",
                                     '#C44E52')

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setWidget(container)
            layout.addWidget(scroll)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка создания графиков: {str(e)}")



    def _create_interactive_bar_chart(self, labels, values, title, color=None):
        """Создание интерактивной столбчатой диаграммы"""
        try:
            values_array = np.array(values, dtype=float)
            if len(labels) != len(values_array):
                raise ValueError("Количество меток и значений не совпадает")

            fig, ax = plt.subplots(figsize=(10, 7))
            fig.subplots_adjust(top=0.9, bottom=0.3, left=0.15, right=0.95)

            # Форматирование значений
            if hasattr(self, 'display_percent') and self.display_percent:
                total = np.sum(values_array)
                if total > 0:
                    values_to_plot = (values_array / total) * 100
                    ylabel = "Приоритет, %"
                    fmt = "{:.1f}%"
                    ylim = (0, 100)
                else:
                    raise ValueError("Невозможно отобразить в процентах: сумма значений равна 0")
            else:
                values_to_plot = values_array
                ylabel = "Значение приоритета"
                fmt = "{:.3f}"
                ylim = (0, np.max(values_array) * 1.15)

            # Создание столбцов
            x_pos = np.arange(len(labels))
            bars = ax.bar(x_pos, values_to_plot,
                          color=color or '#4C72B0',
                          alpha=0.85,
                          width=0.7)

            # Настройка оформления
            ax.set_title(title, pad=25, fontsize=14, fontweight='bold')
            ax.set_ylabel(ylabel, fontsize=12, labelpad=15)
            ax.set_ylim(ylim)
            ax.set_xticks(x_pos)

            formatted_labels = [label.replace(' ', '\n') if len(label) > 10 else label
                                for label in labels]
            ax.set_xticklabels(formatted_labels,
                               fontsize=11,
                               rotation=45,
                               ha='right',
                               rotation_mode='anchor')

            # Подписи значений
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width() / 2.,
                        height + 0.01 * ylim[1],
                        fmt.format(height),
                        ha='center',
                        va='bottom',
                        fontsize=11,
                        fontweight='bold')

            ax.yaxis.grid(True, linestyle='--', alpha=0.6)
            ax.set_axisbelow(True)

            for spine in ['top', 'right']:
                ax.spines[spine].set_visible(False)

            canvas = FigureCanvas(fig)
            canvas.setMinimumSize(800, 500)
            self.res_display_layout.addWidget(canvas)

        except Exception as e:
            error_msg = f"Ошибка при создании столбчатой диаграммы: {str(e)}"
            print(error_msg)
            QMessageBox.warning(self, "Ошибка", error_msg)
            if 'fig' in locals():
                plt.close(fig)


    def _display_table_results(self, layout):
        """Отображение результатов в табличном виде"""
        try:
            if 'priorities' not in self.result_data:
                return

            priorities = self.result_data['priorities']

            # Приоритеты типов критериев (для 3 уровней)
            if self.selected_levels >= 3 and 'type_priority' in priorities:
                types = list(self.backend.criteria_types.keys())
                if len(types) == len(priorities['type_priority']):
                    self._create_priority_table(
                        layout, types, priorities['type_priority'],
                        "Приоритеты видов критериев (Первый уровень)"
                    )

            # Приоритеты критериев (для 2 и 3 уровней)
            if 'criteria_priority' in priorities:
                criteria = self.backend.criteria
                if len(criteria) == len(priorities['criteria_priority']):
                    title = "Приоритеты критериев" + (" (Второй уровень)" if self.selected_levels >= 3 else "")
                    self._create_priority_table(
                        layout, criteria, priorities['criteria_priority'], title
                    )

            # Приоритеты альтернатив
            if 'alternatives_priority' in priorities:
                alts = self.backend.alternatives
                if len(alts) == len(priorities['alternatives_priority']):
                    self._create_priority_table(
                        layout, alts, priorities['alternatives_priority'],
                        "Итоговые приоритеты альтернатив",
                        show_percent=True
                    )

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка создания таблицы: {str(e)}")


    def _display_diagram_results(self, layout):
        """Отображение результатов в виде круговых диаграмм с учетом уровней иерархии"""
        try:
            self._clear_layout(layout)

            container = QWidget()
            container_layout = QVBoxLayout(container)
            container_layout.setSpacing(30)
            container_layout.setContentsMargins(20, 20, 20, 20)

            colors = ['#4C72B0', '#55A868', '#C44E52', '#8172B2', '#CCB974', '#64B5CD']
            textprops = {'fontsize': 14, 'fontweight': 'bold', 'color': '#333333'}
            explode = 0.05

            def autopct_format(values):
                if self.display_percent:
                    return lambda p: f'{p:.1f}%'
                else:
                    total = sum(values)
                    return lambda p: f'{p * total / 100:.3f}\n({p:.1f}%)'

            # Для всех уровней - график альтернатив
            if 'alternatives_priority' in self.result_data['priorities']:
                alts = self.backend.alternatives
                values = self.result_data['priorities']['alternatives_priority']
                if len(alts) == len(values):
                    fig = plt.figure(figsize=(12, 8), facecolor='#f8f8f8')
                    ax = fig.add_subplot(111)

                    wedges, texts, autotexts = ax.pie(
                        values,
                        labels=alts,
                        autopct=autopct_format(values),
                        startangle=90,
                        colors=colors,
                        explode=[explode] * len(alts),
                        shadow={'ox': -0.02, 'oy': 0.02, 'shade': 0.3},
                        textprops=textprops,
                        wedgeprops={'linewidth': 2, 'edgecolor': 'white'},
                        pctdistance=0.75
                    )

                    ax.set_title("ПРИОРИТЕТЫ АЛЬТЕРНАТИВ",
                                 pad=25, fontsize=16, fontweight='bold', color='#2a2a2a')

                    legend = ax.legend(
                        wedges,
                        [f"{a}: {v:.3f}" for a, v in zip(alts, values)],
                        loc='center left',
                        bbox_to_anchor=(1.25, 0.5),
                        fontsize=12,
                        title="Альтернативы",
                        title_fontsize=14,
                        labelspacing=1.5
                    )
                    legend.get_frame().set_facecolor('#f0f0f0')

                    for autotext in autotexts:
                        autotext.set_fontsize(12)
                        autotext.set_fontweight('bold')

                    canvas = FigureCanvas(fig)
                    canvas.setMinimumHeight(600)
                    container_layout.addWidget(canvas, alignment=Qt.AlignTop)

            # Для 2 и 3 уровней - график критериев
            if self.selected_levels >= 2 and 'criteria_priority' in self.result_data['priorities']:
                criteria = self.backend.criteria
                values = self.result_data['priorities']['criteria_priority']
                if len(criteria) == len(values):
                    fig = plt.figure(figsize=(12, 8), facecolor='#f8f8f8')
                    ax = fig.add_subplot(111)

                    wedges, texts, autotexts = ax.pie(
                        values,
                        labels=criteria,
                        autopct=autopct_format(values),
                        startangle=90,
                        colors=colors,
                        explode=[explode] * len(criteria),
                        shadow={'ox': -0.02, 'oy': 0.02, 'shade': 0.3},
                        textprops=textprops,
                        wedgeprops={'linewidth': 2, 'edgecolor': 'white'},
                        pctdistance=0.75
                    )

                    title = "ПРИОРИТЕТЫ КРИТЕРИЕВ" + (" (Второй уровень)" if self.selected_levels >= 3 else "")
                    ax.set_title(title, pad=25, fontsize=16, fontweight='bold', color='#2a2a2a')

                    legend = ax.legend(
                        wedges,
                        [f"{c}: {v:.3f}" for c, v in zip(criteria, values)],
                        loc='center left',
                        bbox_to_anchor=(1.25, 0.5),
                        fontsize=12,
                        title="Критерии",
                        title_fontsize=14,
                        labelspacing=1.5
                    )
                    legend.get_frame().set_facecolor('#f0f0f0')

                    for autotext in autotexts:
                        autotext.set_fontsize(12)
                        autotext.set_fontweight('bold')

                    canvas = FigureCanvas(fig)
                    canvas.setMinimumHeight(600)
                    container_layout.addWidget(canvas, alignment=Qt.AlignTop)

            # Только для 3 уровня - график типов критериев
            if self.selected_levels >= 3 and 'type_priority' in self.result_data['priorities']:
                types = list(self.backend.criteria_types.keys())
                values = self.result_data['priorities']['type_priority']
                if len(types) == len(values):
                    fig = plt.figure(figsize=(12, 8), facecolor='#f8f8f8')
                    ax = fig.add_subplot(111)

                    wedges, texts, autotexts = ax.pie(
                        values,
                        labels=types,
                        autopct=autopct_format(values),
                        startangle=90,
                        colors=colors,
                        explode=[explode] * len(types),
                        shadow={'ox': -0.02, 'oy': 0.02, 'shade': 0.3},
                        textprops=textprops,
                        wedgeprops={'linewidth': 2, 'edgecolor': 'white'},
                        pctdistance=0.75
                    )

                    ax.set_title("ПРИОРИТЕТЫ ВИДОВ КРИТЕРИЕВ (Первый уровень)",
                                 pad=25, fontsize=16, fontweight='bold', color='#2a2a2a')

                    legend = ax.legend(
                        wedges,
                        [f"{t}: {v:.3f}" for t, v in zip(types, values)],
                        loc='center left',
                        bbox_to_anchor=(1.25, 0.5),
                        fontsize=12,
                        title="Типы критериев",
                        title_fontsize=14,
                        labelspacing=1.5
                    )
                    legend.get_frame().set_facecolor('#f0f0f0')
                    legend.get_frame().set_edgecolor('none')

                    for autotext in autotexts:
                        autotext.set_fontsize(12)
                        autotext.set_fontweight('bold')

                    canvas = FigureCanvas(fig)
                    canvas.setMinimumHeight(600)
                    container_layout.addWidget(canvas, alignment=Qt.AlignTop)

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setWidget(container)
            layout.addWidget(scroll)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка создания диаграмм: {str(e)}")

    def _create_table(self, labels, values, title, show_percent=False,
                      is_comparison_matrix=False, is_second_level=False,
                      first_level_weights=None):
        """Создание таблицы с результатами анализа, включая главный вектор"""
        try:
            # Проверка входных данных
            if not labels or values is None:
                raise ValueError("Не переданы labels или values")
            if is_second_level and first_level_weights is None:
                raise ValueError("Для таблицы второго уровня нужны first_level_weights")

            # Создаем группу для таблицы
            group = QGroupBox(title)
            group.setStyleSheet("""
                QGroupBox {
                    border: 1px solid #e0e0e0;
                    border-radius: 4px;
                    margin-top: 10px;
                    padding-top: 15px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 5px;
                }
            """)

            layout = QVBoxLayout(group)
            layout.setContentsMargins(5, 15, 5, 10)
            layout.setSpacing(10)

            # Создаем таблицу с нужным количеством столбцов
            table = QTableWidget()
            columns = 4  # Убрали столбец с процентами/долей, оставили только главный вектор и вектор приоритетов
            table.setColumnCount(columns)

            headers = ["№", "Элемент", "Главный вектор (ΓB)", "Вектор приоритетов (w)"]
            table.setHorizontalHeaderLabels(headers)

            # Настройка таблицы
            table.verticalHeader().setVisible(False)
            table.setEditTriggers(QTableWidget.NoEditTriggers)
            table.setSelectionMode(QTableWidget.NoSelection)
            table.setMinimumHeight(min(300, len(labels) * 35 + 40))

            # Настройка размеров столбцов
            for i in range(columns):
                resize_mode = QHeaderView.ResizeToContents if i != 1 else QHeaderView.Stretch
                table.horizontalHeader().setSectionResizeMode(i, resize_mode)

            # Обработка значений и расчет векторов
            values_array = np.array(values, dtype=float)

            if values_array.ndim == 1:
                # Если переданы готовые приоритеты (w)
                w = values_array
                # Рассчитываем главный вектор ΓB из вектора приоритетов
                n = len(w)
                CB = np.array([np.prod([(w[i] / w[j]) if w[j] != 0 else 0
                                        for j in range(n)]) ** (1 / n) for i in range(n)])

                # Нормализуем главный вектор
                CB = CB / np.sum(CB) if np.sum(CB) > 0 else CB
                w_normalized = w / np.sum(w) if np.sum(w) > 0 else w
            else:
                # Если передана матрица сравнения
                CB, w = self.backend.calculate_priority_vector(values_array)
                w_normalized = w / np.sum(w) if np.sum(w) > 0 else w

            # Заполнение таблицы данными
            table.setRowCount(len(labels))
            max_w = np.max(w_normalized) if len(w_normalized) > 0 else 0

            for row in range(len(labels)):
                # Создаем элементы таблицы
                items = [
                    QTableWidgetItem(str(row + 1)),  # №
                    QTableWidgetItem(str(labels[row])),  # Элемент
                    QTableWidgetItem(f"{CB[row]:.6f}"),  # Главный вектор ΓB
                    QTableWidgetItem(f"{w_normalized[row]:.6f}")  # Вектор приоритетов
                ]

                # Настройка стилей
                alignments = [
                    Qt.AlignCenter,
                    Qt.AlignLeft | Qt.AlignVCenter,
                    Qt.AlignRight | Qt.AlignVCenter,
                    Qt.AlignRight | Qt.AlignVCenter
                ]

                # Подсветка максимального значения
                if w_normalized[row] == max_w:
                    highlight_color = QColor(235, 245, 235)
                    text_color = QColor(0, 100, 0)
                    font = QFont()
                    font.setBold(True)
                else:
                    highlight_color = QColor(255, 255, 255)
                    text_color = QColor(0, 0, 0)
                    font = QFont()

                # Устанавливаем свойства для всех ячеек строки
                for col in range(len(items)):
                    items[col].setTextAlignment(alignments[col])
                    items[col].setBackground(highlight_color)
                    items[col].setForeground(text_color)
                    items[col].setFont(font)
                    table.setItem(row, col, items[col])

            # Добавление информации о согласованности для матриц сравнения
            if is_comparison_matrix and values_array.ndim > 1 and values_array.shape[0] > 2:
                consistency = self.backend.check_consistency(values_array)
                status_text = (
                    f"Согласованность: λmax = {consistency['lambda_max']:.3f}, "
                    f"ИС = {consistency['CI']:.3f}, ОС = {consistency['CR']:.3f} - "
                    f"{consistency['status']}"
                )
                status_label = QLabel(status_text)

                if consistency['CR'] < 0.1:
                    status_label.setStyleSheet("color: green;")
                elif consistency['CR'] < 0.2:
                    status_label.setStyleSheet("color: orange;")
                else:
                    status_label.setStyleSheet("color: red; font-weight: bold;")

                layout.addWidget(status_label)

            layout.addWidget(table)
            self.res_display_layout.addWidget(group)
            return True

        except Exception as e:
            error_msg = f"Не удалось создать таблицу:\n{str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Ошибка", error_msg)
            return False

    def _add_consistency_info(self, layout, matrix):
        """Добавляет информацию о согласованности матрицы"""
        consistency = self.backend.check_consistency(matrix)
        status_text = (
            f"Согласованность: λmax = {consistency['lambda_max']:.3f}, "
            f"ИС = {consistency['CI']:.3f}, ОС = {consistency['CR']:.3f} - "
            f"{consistency['status']}"
        )
        status_label = QLabel(status_text)

        if consistency['CR'] < 0.1:
            status_label.setStyleSheet("color: green;")
        elif consistency['CR'] < 0.2:
            status_label.setStyleSheet("color: orange;")
        else:
            status_label.setStyleSheet("color: red; font-weight: bold;")

        layout.addWidget(status_label)

    def _add_table_row(self, table, row_num, label, CB_value, w_value, max_w,
                       show_percent, is_second_level, second_level_w=None):
        """Добавляет строку в таблицу результатов"""
        # Создаем элементы таблицы
        items = [
            QTableWidgetItem(str(row_num + 1)),
            QTableWidgetItem(str(label)),
            QTableWidgetItem(f"{CB_value:.6f}"),
            QTableWidgetItem(f"{w_value * 100:.2f}%" if show_percent else f"{w_value:.6f}")
        ]

        if is_second_level and second_level_w is not None:
            items.append(QTableWidgetItem(
                f"{second_level_w[row_num] * 100:.2f}%" if show_percent else f"{second_level_w[row_num]:.6f}"
            ))

        # Настройка стилей и выравнивания
        alignments = [Qt.AlignCenter,
                      Qt.AlignLeft | Qt.AlignVCenter,
                      Qt.AlignRight | Qt.AlignVCenter,
                      Qt.AlignRight | Qt.AlignVCenter]

        if is_second_level:
            alignments.append(Qt.AlignRight | Qt.AlignVCenter)

        # Подсветка максимального значения
        is_max = w_value == max_w
        highlight_color = QColor(235, 245, 235) if is_max else QColor(255, 255, 255)
        text_color = QColor(0, 100, 0) if is_max else QColor(0, 0, 0)

        font = QFont()
        if is_max:
            font.setBold(True)

        # Устанавливаем свойства для всех ячеек строки
        for col in range(len(items)):
            items[col].setTextAlignment(alignments[col])
            items[col].setBackground(highlight_color)
            items[col].setForeground(text_color)
            items[col].setFont(font)
            table.setItem(row_num, col, items[col])


    def _get_color_for_value(self, value, min_val, max_val):
        """Получение цвета для значения"""
        try:
            normalized = (value - min_val) / (max_val - min_val)
            red = min(255, int(255 * (1 - normalized * 2) if normalized < 0.5 else 0))
            green = min(255, int(255 * normalized * 2 if normalized < 0.5 else 255))
            blue = 0
            return QColor(red, green, blue)
        except:
            return QColor(255, 255, 255)


    def _create_bar_chart(self, labels, values, title, show_percent=False):
        """Создание столбчатой диаграммы"""
        try:
            fig, ax = plt.subplots(figsize=(10, 4))

            if show_percent:
                total = np.sum(values)
                if total > 0:
                    values = [v / total * 100 for v in values]
                ylabel = "Приоритет, %"
            else:
                ylabel = "Значение приоритета"

            bars = ax.bar(labels, values)
            ax.set_title(title)
            ax.set_ylabel(ylabel)
            plt.xticks(rotation=45, ha='right')

            for bar in bars:
                height = bar.get_height()
                ax.text(
                    bar.get_x() + bar.get_width() / 2.,
                    height,
                    f'{height:.1f}%' if show_percent else f'{height:.3f}',
                    ha='center', va='bottom'
                )

            canvas = FigureCanvas(fig)
            self.res_display_layout.addWidget(canvas)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка создания столбчатой диаграммы: {str(e)}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    try:
        window = AHPFrontend()
        window.show()
        sys.exit(app.exec_())
    except Exception as e:
        QMessageBox.critical(None, "Критическая ошибка", f"Программа завершилась с ошибкой:\n{str(e)}")
        sys.exit(1)
